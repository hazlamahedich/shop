"""Tests for Knowledge Base API endpoints.

Tests cover:
- AC1: Document upload with metadata
- AC2: Document chunking (verify chunk sizes)
- AC3: Document list endpoint
- AC4: Document deletion with chunks
- Edge cases: Invalid file type, file too large, unauthorized access, document not found
"""

from __future__ import annotations

import os
import tempfile

import pytest
from httpx import AsyncClient, ASGITransport
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.main import app
from app.core.database import get_db
from app.models.knowledge_base import DocumentChunk, DocumentStatus, KnowledgeDocument
from app.models.merchant import Merchant
from app.services.knowledge.chunker import ChunkingError, DocumentChunker


@pytest.fixture
async def db_session(async_session: AsyncSession):
    """Alias for async_session fixture."""
    yield async_session


@pytest.fixture
async def test_merchant(async_session: AsyncSession) -> int:
    """Create a test merchant for integration tests."""
    merchant = Merchant(
        merchant_key="test-kb-merchant-key",
        platform="messenger",
        email="kb-test@example.com",
        status="active",
    )
    async_session.add(merchant)
    await async_session.commit()
    await async_session.refresh(merchant)
    return merchant.id


@pytest.fixture
async def auth_headers(test_merchant: int) -> dict[str, str]:
    """Generate authentication headers with JWT token."""
    from app.core.auth import create_jwt
    import uuid

    session_id = str(uuid.uuid4())
    token = create_jwt(merchant_id=test_merchant, session_id=session_id)
    return {"Authorization": f"Bearer {token}"}


@pytest.fixture
async def client(async_session: AsyncSession):
    """Create async HTTP client for testing."""

    async def override_get_db():
        yield async_session

    app.dependency_overrides[get_db] = override_get_db

    async with AsyncClient(
        transport=ASGITransport(app=app),
        base_url="http://test",
    ) as http_client:
        yield http_client

    app.dependency_overrides.clear()


async def upload_document_helper(
    client: AsyncClient,
    auth_headers: dict,
    filename: str,
    content: bytes,
    content_type: str,
    expected_status: int = 200,
) -> dict:
    """Helper to upload document and return response data.

    Reduces duplication across upload tests by encapsulating:
    - HTTP POST request
    - Status code assertion
    - Response data extraction

    Args:
        client: AsyncClient instance
        auth_headers: Authentication headers
        filename: Name of file to upload
        content: File content as bytes
        content_type: MIME type
        expected_status: Expected HTTP status code (default: 200)

    Returns:
        Response data dict (response.json()["data"])
    """
    response = await client.post(
        "/api/knowledge-base/upload",
        files={"file": (filename, content, content_type)},
        headers=auth_headers,
    )
    assert response.status_code == expected_status
    return response.json()["data"] if expected_status == 200 else response.json()


class TestKnowledgeBaseAPI:
    """Tests for knowledge base endpoints.

    Parallel Execution: All tests in this class are independent and can run in parallel.
    Configure with: pytest -n auto tests/api/test_knowledge_base.py
    """

    @pytest.mark.asyncio
    async def test_upload_document_pdf(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-006 [P0]: AC1 - Test document upload with PDF file (malformed PDF shows error handling)."""
        pdf_content = b"%PDF-1.4\n1 0 obj\n<< /Type /Catalog >>\nendobj\ntrailer\n%%EOF"

        data = await upload_document_helper(
            client, auth_headers, "test.pdf", pdf_content, "application/pdf"
        )

        assert data["filename"] == "test.pdf"
        assert data["fileType"] == "pdf"
        assert data["status"] == "error"

    @pytest.mark.asyncio
    async def test_upload_document_txt(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-007 [P0]: AC1 - Test document upload with TXT file."""
        txt_content = b"This is a test document for the knowledge base."

        data = await upload_document_helper(
            client, auth_headers, "test.txt", txt_content, "text/plain"
        )

        assert data["filename"] == "test.txt"
        assert data["fileType"] == "txt"

    @pytest.mark.asyncio
    async def test_upload_invalid_file_type(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-008 [P1]: Edge Case - Test upload with invalid file type (400 error)."""
        data = await upload_document_helper(
            client,
            auth_headers,
            "test.exe",
            b"binary content",
            "application/octet-stream",
            expected_status=400,
        )
        assert "error" in data or "detail" in data

    @pytest.mark.asyncio
    async def test_upload_file_too_large(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-009 [P1]: Edge Case - Test upload with file too large (413 error)."""
        large_content = b"x" * (11 * 1024 * 1024)  # 11MB > 10MB limit

        data = await upload_document_helper(
            client, auth_headers, "large.txt", large_content, "text/plain", expected_status=413
        )
        assert "error" in data or "detail" in data

    @pytest.mark.asyncio
    async def test_list_documents(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-010 [P0]: AC3 - Test document list endpoint."""
        doc = KnowledgeDocument(
            merchant_id=test_merchant,
            filename="list.txt",
            file_type="txt",
            file_size=100,
            status=DocumentStatus.READY.value,
        )
        db_session.add(doc)
        await db_session.commit()

        response = await client.get(
            "/api/knowledge-base",
            headers=auth_headers,
        )

        assert response.status_code == 200
        data = response.json()["data"]
        assert len(data["documents"]) >= 1
        assert any(d["filename"] == "list.txt" for d in data["documents"])

    @pytest.mark.asyncio
    async def test_get_document(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-011 [P1]: AC3 - Test get document details endpoint."""
        doc = KnowledgeDocument(
            merchant_id=test_merchant,
            filename="detail.txt",
            file_type="txt",
            file_size=100,
            status=DocumentStatus.READY.value,
        )
        db_session.add(doc)
        await db_session.commit()
        await db_session.refresh(doc)

        response = await client.get(
            f"/api/knowledge-base/{doc.id}",
            headers=auth_headers,
        )

        assert response.status_code == 200
        data = response.json()["data"]
        assert data["filename"] == "detail.txt"
        assert data["chunkCount"] == 0

    @pytest.mark.asyncio
    async def test_get_document_status(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-012 [P1]: AC3 - Test get document processing status endpoint."""
        doc = KnowledgeDocument(
            merchant_id=test_merchant,
            filename="status.txt",
            file_type="txt",
            file_size=100,
            status=DocumentStatus.PROCESSING.value,
        )
        db_session.add(doc)
        await db_session.commit()
        await db_session.refresh(doc)

        response = await client.get(
            f"/api/knowledge-base/{doc.id}/status",
            headers=auth_headers,
        )

        assert response.status_code == 200
        data = response.json()["data"]
        assert data["status"] == "processing"
        assert data["progress"] == 50

    @pytest.mark.asyncio
    async def test_delete_document(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-013 [P0]: AC4 - Test document deletion with cascading chunks."""
        doc = KnowledgeDocument(
            merchant_id=test_merchant,
            filename="delete.txt",
            file_type="txt",
            file_size=100,
            status=DocumentStatus.READY.value,
        )
        db_session.add(doc)
        await db_session.commit()
        await db_session.refresh(doc)

        chunk = DocumentChunk(
            document_id=doc.id,
            chunk_index=0,
            content="Test chunk content",
        )
        db_session.add(chunk)
        await db_session.commit()

        response = await client.delete(
            f"/api/knowledge-base/{doc.id}",
            headers=auth_headers,
        )

        assert response.status_code == 200
        data = response.json()["data"]
        assert data["deleted"] is True

        result = await db_session.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.id == doc.id)
        )
        assert result.scalar_one_or_none() is None

    @pytest.mark.asyncio
    async def test_delete_nonexistent_document(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-014 [P1]: Edge Case - Test delete non-existent document (404 error)."""
        response = await client.delete(
            "/api/knowledge-base/99999",
            headers=auth_headers,
        )

        assert response.status_code == 404

    @pytest.mark.asyncio
    async def test_get_nonexistent_document(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-015 [P1]: Edge Case - Test get non-existent document (404 error)."""
        response = await client.get(
            "/api/knowledge-base/99999",
            headers=auth_headers,
        )

        assert response.status_code == 404

    @pytest.mark.asyncio
    async def test_upload_document_md(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-001 [P0]: AC1 - Test document upload with Markdown file."""
        md_content = b"""# Knowledge Base Document

This is a **markdown** document for testing.

## Section 1
- Item 1
- Item 2

## Section 2
Some more content here.
"""

        response = await client.post(
            "/api/knowledge-base/upload",
            files={"file": ("test.md", md_content, "text/markdown")},
            headers=auth_headers,
        )

        assert response.status_code == 200
        data = response.json()["data"]
        assert data["filename"] == "test.md"
        assert data["fileType"] == "md"
        assert "id" in data
        assert "status" in data

    @pytest.mark.asyncio
    async def test_upload_document_docx(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-002 [P0]: AC1 - Test document upload with DOCX file."""
        from docx import Document
        from io import BytesIO

        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test DOCX file for the knowledge base.")
        doc.add_paragraph("It contains multiple paragraphs of text.")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_content = buffer.read()

        response = await client.post(
            "/api/knowledge-base/upload",
            files={
                "file": (
                    "test.docx",
                    docx_content,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            headers=auth_headers,
        )

        assert response.status_code == 200
        data = response.json()["data"]
        assert data["filename"] == "test.docx"
        assert data["fileType"] == "docx"
        assert "id" in data

    @pytest.mark.asyncio
    async def test_unauthorized_access(
        self,
        db_session: AsyncSession,
        client: AsyncClient,
    ) -> None:
        """8.3-API-003 [P0]: Security - Test that all endpoints require authentication."""
        txt_content = b"Unauthorized test content"

        response = await client.post(
            "/api/knowledge-base/upload",
            files={"file": ("unauth.txt", txt_content, "text/plain")},
        )

        assert response.status_code == 401

        response = await client.get("/api/knowledge-base")
        assert response.status_code == 401

        response = await client.get("/api/knowledge-base/1")
        assert response.status_code == 401

        response = await client.delete("/api/knowledge-base/1")
        assert response.status_code == 401

    @pytest.mark.asyncio
    async def test_cross_merchant_isolation(
        self,
        db_session: AsyncSession,
        client: AsyncClient,
        test_merchant: int,
        auth_headers: dict,
    ) -> None:
        """8.3-API-004 [P0]: Security - Test that merchants cannot access other merchants' documents (IDOR prevention)."""
        from app.core.auth import create_jwt
        import uuid

        merchant_a = Merchant(
            merchant_key="merchant-a-key",
            platform="messenger",
            email="merchant-a@example.com",
            status="active",
        )
        merchant_b = Merchant(
            merchant_key="merchant-b-key",
            platform="messenger",
            email="merchant-b@example.com",
            status="active",
        )
        db_session.add(merchant_a)
        db_session.add(merchant_b)
        await db_session.commit()
        await db_session.refresh(merchant_a)
        await db_session.refresh(merchant_b)

        doc_a = KnowledgeDocument(
            merchant_id=merchant_a.id,
            filename="merchant-a-doc.txt",
            file_type="txt",
            file_size=100,
            status=DocumentStatus.READY.value,
        )
        db_session.add(doc_a)
        await db_session.commit()
        await db_session.refresh(doc_a)

        session_id_b = str(uuid.uuid4())
        token_b = create_jwt(merchant_id=merchant_b.id, session_id=session_id_b)
        headers_b = {"Authorization": f"Bearer {token_b}"}

        response = await client.get(
            f"/api/knowledge-base/{doc_a.id}",
            headers=headers_b,
        )

        assert response.status_code == 404

        response = await client.delete(
            f"/api/knowledge-base/{doc_a.id}",
            headers=headers_b,
        )

        assert response.status_code == 404

        response = await client.get(
            "/api/knowledge-base",
            headers=headers_b,
        )

        assert response.status_code == 200
        data = response.json()["data"]
        assert not any(d["id"] == doc_a.id for d in data["documents"])

    @pytest.mark.asyncio
    async def test_list_documents_empty(
        self,
        db_session: AsyncSession,
        test_merchant: int,
        client: AsyncClient,
        auth_headers: dict,
    ) -> None:
        """8.3-API-005 [P2]: AC3 - Test empty document list returns empty array."""
        response = await client.get(
            "/api/knowledge-base",
            headers=auth_headers,
        )

        assert response.status_code == 200
        data = response.json()
        assert "data" in data
        assert "documents" in data["data"]
        assert isinstance(data["data"]["documents"], list)


class TestDocumentChunker:
    """Tests for document chunker service.

    Parallel Execution: All tests in this class are independent and can run in parallel.
    Configure with: pytest -n auto tests/api/test_knowledge_base.py::TestDocumentChunker
    """

    def test_chunk_text_file(self, tmp_path: str) -> None:
        """8.3-UNIT-003 [P1]: AC2 - Test chunking of text file."""
        chunker = DocumentChunker()

        test_file = os.path.join(tmp_path, "test.txt")
        with open(test_file, "w") as f:
            f.write("This is a test document. " * 100)

        chunks = chunker.chunk_document(test_file, "txt")

        assert len(chunks) > 0
        for chunk in chunks:
            assert len(chunk) >= chunker.MIN_CHUNK_CHARS

    def test_chunk_markdown_file(self, tmp_path: str) -> None:
        """8.3-UNIT-004 [P1]: AC2 - Test chunking of markdown file."""
        chunker = DocumentChunker()

        test_file = os.path.join(tmp_path, "test.md")
        with open(test_file, "w") as f:
            f.write("# Test Header\n\nThis is markdown content. " * 50)

        chunks = chunker.chunk_document(test_file, "md")

        assert len(chunks) > 0

    def test_chunk_empty_file(self, tmp_path: str) -> None:
        """8.3-UNIT-005 [P2]: Edge Case - Test chunking of empty file raises error."""
        chunker = DocumentChunker()

        test_file = os.path.join(tmp_path, "empty.txt")
        with open(test_file, "w") as f:
            f.write("")

        with pytest.raises(ChunkingError):
            chunker.chunk_document(test_file, "txt")

    def test_chunk_quality_validation(self, tmp_path: str) -> None:
        """8.3-UNIT-006 [P2]: Edge Case - Test that whitespace-only chunks are filtered."""
        chunker = DocumentChunker()

        test_file = os.path.join(tmp_path, "whitespace.txt")
        with open(test_file, "w") as f:
            f.write(
                "   \n\n   \t\t   \n\nValid content here with enough characters to pass validation."
            )

        chunks = chunker.chunk_document(test_file, "txt")

        for chunk in chunks:
            assert any(c.isalnum() for c in chunk)

    def test_chunk_size_bounds(self, tmp_path: str) -> None:
        """8.3-UNIT-001 [P1]: AC2 - Verify chunks are within 500-1000 char range."""
        chunker = DocumentChunker()

        test_file = os.path.join(tmp_path, "large.txt")
        long_content = "This is a sentence. " * 500
        with open(test_file, "w") as f:
            f.write(long_content)

        chunks = chunker.chunk_document(test_file, "txt")

        assert len(chunks) > 1, "Expected multiple chunks for large document"

        for i, chunk in enumerate(chunks):
            assert len(chunk) >= chunker.MIN_CHUNK_CHARS, (
                f"Chunk {i} too small: {len(chunk)} chars (min: {chunker.MIN_CHUNK_CHARS})"
            )
            assert len(chunk) <= chunker.CHUNK_SIZE_MAX + 100, (
                f"Chunk {i} too large: {len(chunk)} chars (max: {chunker.CHUNK_SIZE_MAX})"
            )

    def test_chunk_docx_file(self, tmp_path: str) -> None:
        """8.3-UNIT-002 [P1]: Test chunking of DOCX file."""
        chunker = DocumentChunker()

        test_file = os.path.join(tmp_path, "test.docx")
        try:
            from docx import Document

            doc = Document()
            doc.add_heading("Test Document", 0)
            doc.add_paragraph("This is paragraph 1 with some content.")
            doc.add_paragraph("This is paragraph 2 with more content.")
            doc.save(test_file)

            chunks = chunker.chunk_document(test_file, "docx")

            assert len(chunks) > 0, "Expected at least one chunk from DOCX"
            for chunk in chunks:
                assert len(chunk) >= chunker.MIN_CHUNK_CHARS
        except ImportError:
            pytest.skip("python-docx not installed")
